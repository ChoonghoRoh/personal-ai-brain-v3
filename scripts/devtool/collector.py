#!/usr/bin/env python3
"""
다양한 문서 포맷 수집 및 처리 시스템
PDF, DOCX 파일 지원
"""

import sys
from pathlib import Path
from typing import List, Dict, Optional

# 프로젝트 루트 경로 설정
PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent
BRAIN_DIR = PROJECT_ROOT / "brain"
COLLECTOR_DIR = PROJECT_ROOT / "collector"

try:
    from pypdf import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def extract_text_from_pdf(pdf_path: Path) -> Optional[str]:
    """PDF 파일에서 텍스트 추출"""
    if not PDF_AVAILABLE:
        print(f"⚠️  pypdf가 설치되지 않았습니다. PDF 처리 불가: {pdf_path}")
        return None
    
    try:
        reader = PdfReader(pdf_path)
        text_parts = []
        
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        
        return "\n\n".join(text_parts)
    except Exception as e:
        print(f"❌ PDF 읽기 오류 {pdf_path}: {e}")
        return None


def extract_text_from_docx(docx_path: Path) -> Optional[str]:
    """DOCX 파일에서 텍스트 추출"""
    if not DOCX_AVAILABLE:
        print(f"⚠️  python-docx가 설치되지 않았습니다. DOCX 처리 불가: {docx_path}")
        return None
    
    try:
        doc = Document(docx_path)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # 테이블에서도 텍스트 추출
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return "\n\n".join(text_parts)
    except Exception as e:
        print(f"❌ DOCX 읽기 오류 {docx_path}: {e}")
        return None


def convert_to_markdown(source_file: Path, output_dir: Path) -> Optional[Path]:
    """문서를 Markdown으로 변환하여 저장"""
    if not source_file.exists():
        return None
    
    # 텍스트 추출
    text = None
    if source_file.suffix.lower() == '.pdf':
        text = extract_text_from_pdf(source_file)
    elif source_file.suffix.lower() in ['.docx', '.doc']:
        text = extract_text_from_docx(source_file)
    else:
        print(f"⚠️  지원하지 않는 파일 형식: {source_file.suffix}")
        return None
    
    if not text or not text.strip():
        print(f"⚠️  텍스트를 추출할 수 없습니다: {source_file}")
        return None
    
    # 출력 파일 경로 생성
    output_dir.mkdir(parents=True, exist_ok=True)
    output_file = output_dir / f"{source_file.stem}.md"
    
    # Markdown 파일로 저장
    try:
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(f"# {source_file.stem}\n\n")
            f.write(f"*원본 파일: {source_file.name}*\n\n")
            f.write("---\n\n")
            f.write(text)
        
        print(f"✅ 변환 완료: {source_file.name} → {output_file.name}")
        return output_file
    except Exception as e:
        print(f"❌ 파일 저장 오류: {e}")
        return None


def process_collector_directory():
    """collector 디렉토리의 파일들을 처리"""
    if not COLLECTOR_DIR.exists():
        COLLECTOR_DIR.mkdir(parents=True, exist_ok=True)
        print(f"📁 collector 디렉토리 생성: {COLLECTOR_DIR}")
        return []
    
    converted_files = []
    
    # PDF 파일 처리
    for pdf_file in COLLECTOR_DIR.rglob("*.pdf"):
        if pdf_file.is_file():
            md_file = convert_to_markdown(pdf_file, BRAIN_DIR / "reference")
            if md_file:
                converted_files.append(md_file)
    
    # DOCX 파일 처리
    for docx_file in COLLECTOR_DIR.rglob("*.docx"):
        if docx_file.is_file():
            md_file = convert_to_markdown(docx_file, BRAIN_DIR / "reference")
            if md_file:
                converted_files.append(md_file)
    
    return converted_files


def main():
    """메인 함수"""
    print("=" * 60)
    print("문서 수집 및 변환 시스템")
    print("=" * 60)
    
    if not PDF_AVAILABLE and not DOCX_AVAILABLE:
        print("❌ PDF 또는 DOCX 처리를 위한 라이브러리가 설치되지 않았습니다.")
        print("   다음 명령어로 설치하세요:")
        print("   pip install pypdf python-docx")
        return
    
    print(f"\n[1/2] collector 디렉토리 스캔 중...")
    print(f"      경로: {COLLECTOR_DIR}")
    
    converted_files = process_collector_directory()
    
    if not converted_files:
        print("\n✅ 처리할 파일이 없습니다.")
        print(f"   PDF/DOCX 파일을 {COLLECTOR_DIR}에 넣고 다시 실행하세요.")
        return
    
    print(f"\n[2/2] 변환 완료: {len(converted_files)}개 파일")
    print(f"\n✅ 변환된 파일은 brain/reference/에 저장되었습니다.")
    print(f"   다음 단계: python scripts/embed_and_store.py 실행하여 임베딩 저장")


if __name__ == "__main__":
    main()

